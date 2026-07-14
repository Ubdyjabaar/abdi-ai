import os
import base64
import logging
import tempfile
from groq import Groq
from docx import Document

logger = logging.getLogger(__name__)

client = Groq(api_key=os.getenv("GROQ_API_KEY"))

def extract_text_from_pdf(file_path: str) -> str:
    import fitz
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text()
    doc.close()
    return text

def encode_image(file_path: str) -> str:
    with open(file_path, "rb") as f:
        return base64.b64encode(f.read()).decode("utf-8")

def solve_quiz_text(content: str) -> str:
    prompt = f"""Solve this quiz/exam. For each question:
1. Write the question
2. Provide the correct answer
3. Give a brief explanation

Quiz content:
{content}

Format the answer clearly with question numbers."""
    
    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.1,
    )
    return response.choices[0].message.content

def solve_quiz_image(file_path: str) -> str:
    image_data = encode_image(file_path)
    response = client.chat.completions.create(
        model="llama-3.2-90b-vision-preview",
        messages=[{
            "role": "user",
            "content": [
                {"type": "text", "text": "Solve this quiz/exam shown in the image. List each question, give the correct answer, and explain briefly."},
                {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image_data}"}}
            ]
        }],
        temperature=0.1,
    )
    return response.choices[0].message.content

def create_solved_document(original_name: str, solution_text: str) -> str:
    doc = Document()
    doc.add_heading(f"Solved: {original_name}", 0)
    doc.add_paragraph("")
    for line in solution_text.split("\n"):
        if line.strip():
            doc.add_paragraph(line)
    
    output_path = os.path.join(tempfile.gettempdir(), f"solved_{os.path.splitext(original_name)[0]}.docx")
    doc.save(output_path)
    return output_path

async def solve_quiz_file(file_path: str, file_name: str) -> str:
    logger.info(f"Solving quiz: {file_name}")
    ext = file_name.lower()
    
    if ext.endswith(".pdf"):
        content = extract_text_from_pdf(file_path)
        if not content.strip():
            raise ValueError("Could not extract text from PDF")
        solution = solve_quiz_text(content[:10000])
    elif ext.endswith((".png", ".jpg", ".jpeg", ".gif", ".webp")):
        solution = solve_quiz_image(file_path)
    else:
        raise ValueError("Unsupported format. Send PDF or image.")
    
    output_path = create_solved_document(file_name, solution)
    return output_path
